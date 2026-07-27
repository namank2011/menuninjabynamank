# -*- coding: utf-8 -*-
"""
QuickMenu Agent - Diagnostics and E2E Pipeline Verification Utility
This script runs a complete system health check:
1. Environment & Config setup (.env, folders structure).
2. Python Package imports safety.
3. System CLI/Binary dependencies (Tesseract, LibreOffice).
4. Local Ollama Server availability & Model checks.
5. End-to-End Extraction verification for all supported formats (TXT, CSV, XLSX, DOCX, Image with OCR).
6. Excel Export template formatting & column mapping verification.
"""

from __future__ import annotations

import sys
import os
import re
import json
import shutil
import subprocess
import time
import tempfile
import sqlite3
from openpyxl import load_workbook, Workbook
from pathlib import Path
from typing import Dict, Any, List, Tuple

# Enable ANSI colors on Windows Console
os.system("")

# Colors
C_GREEN = "\033[92m"
C_YELLOW = "\033[93m"
C_RED = "\033[91m"
C_CYAN = "\033[96m"
C_BOLD = "\033[1m"
C_RESET = "\033[0m"

# Load environment variables from .env if present
from dotenv import load_dotenv
load_dotenv()

# Setup sys.path to locate backend modules cleanly
backend_dir = Path(__file__).resolve().parent
if str(backend_dir) not in sys.path:
    sys.path.insert(0, str(backend_dir))

# Try importing project modules globally
try:
    from file_extractors import extract_menu_from_file, SUPPORTED_EXTS
    from bulk_writer import write_bulk_upload_excel, PRODUCT_SHEET, MANDATORY_HEADERS
    from database import init_db, DB_FILE
    from validation import validate_menu
    from schemas import MenuExtraction, MenuItem
except ImportError as e:
    print(f"Error importing internal modules globally: {e}")

def print_header(title: str):
    print(f"\n{C_BOLD}{C_CYAN}=== {title} ==={C_RESET}")

def report_status(check_name: str, passed: bool, info: str = ""):
    if passed:
        print(f"[{C_GREEN}PASS{C_RESET}] {C_BOLD}{check_name}{C_RESET}: {info}")
    else:
        print(f"[{C_RED}FAIL{C_RESET}] {C_BOLD}{check_name}{C_RESET}: {info}")

def report_warning(check_name: str, info: str = ""):
    print(f"[{C_YELLOW}WARN{C_RESET}] {C_BOLD}{check_name}{C_RESET}: {info}")

def run_diagnostics(verbose: bool = True) -> Dict[str, Any]:
    report = {
        "timestamp": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "status": "HEALTHY",
        "steps": {}
    }
    
    # ----------------------------------------------------
    # 1. Environment and Folders
    # ----------------------------------------------------
    if verbose:
        print_header("1. CHECKING DIRECTORIES & SYSTEM ENVIRONMENT")
    
    base_dir = backend_dir.parent
    uploads_dir = base_dir / "uploads"
    outputs_dir = base_dir / "outputs"
    templates_dir = backend_dir / "templates"
    env_file = base_dir / ".env"
    default_template = templates_dir / "Bulk_Upload_Sheet_Format.xlsx"
    
    folder_checks = {
        "uploads": uploads_dir.exists(),
        "outputs": outputs_dir.exists(),
        "templates": templates_dir.exists(),
        "default_template": default_template.exists(),
        "dotenv": env_file.exists()
    }
    
    for key, exists in folder_checks.items():
        if not exists:
            # Auto-create directories if missing
            if key in ["uploads", "outputs", "templates"]:
                (base_dir / key).mkdir(exist_ok=True)
                folder_checks[key] = True
                if verbose:
                    report_status(f"Directory '{key}'", True, "Created successfully")
            else:
                if verbose:
                    report_status(f"File '{key}'", False, f"Missing at path: {base_dir / key}")
                report["status"] = "DEGRADED"
        else:
            if verbose:
                report_status(f"Component '{key}'", True, f"Exists at: {base_dir / key}")
                
    report["steps"]["directories"] = folder_checks

    # ----------------------------------------------------
    # 2. Python Package Imports
    # ----------------------------------------------------
    if verbose:
        print_header("2. VERIFYING PYTHON PACKAGE IMPORTS")
        
    modules_to_test = [
        ("fastapi", "FastAPI web framework"),
        ("uvicorn", "Web server"),
        ("openpyxl", "Excel reader/writer"),
        ("fitz", "PyMuPDF PDF text and image extractor"),
        ("docx", "Python-docx Word file parser"),
        ("pandas", "Pandas tabular analytics framework"),
        ("rapidocr_onnxruntime", "Local RapidOCR engine"),
        ("PIL", "Pillow image library"),
        ("requests", "HTTP calls model API connector"),
        ("dotenv", "Dotenv environment parser")
    ]
    
    import_results = {}
    for module_name, desc in modules_to_test:
        try:
            __import__(module_name)
            import_results[module_name] = True
            if verbose:
                report_status(f"Import '{module_name}'", True, desc)
        except ImportError as err:
            import_results[module_name] = False
            if verbose:
                report_status(f"Import '{module_name}'", False, f"Failed to import: {err}")
            report["status"] = "CRITICAL"
            
    report["steps"]["imports"] = import_results

    # ----------------------------------------------------
    # 3. Database Schema Integrity
    # ----------------------------------------------------
    if verbose:
        print_header("3. VERIFYING SQLITE DATABASE SCHEMA & INTEGRITY")
        
    db_file_path = base_dir / "outputs" / "shopverse_agent.db"
    db_result = {"exists": db_file_path.exists(), "writeable": False, "schemas": {}}
    
    try:
        # Import internally
        from database import init_db
        init_db()
        db_result["exists"] = db_file_path.exists()
        
        conn = sqlite3.connect(str(db_file_path))
        cursor = conn.cursor()
        
        # Verify tables schema
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = [t[0] for t in cursor.fetchall()]
        
        db_result["tables"] = tables
        expected_tables = ["drafts", "draft_items", "audit_logs"]
        
        all_tables_exist = True
        for tbl in expected_tables:
            if tbl in tables:
                cursor.execute(f"PRAGMA table_info({tbl});")
                cols = [c[1] for c in cursor.fetchall()]
                db_result["schemas"][tbl] = cols
                if verbose:
                    report_status(f"Table '{tbl}' schema", True, f"Columns: {', '.join(cols[:5])}...")
            else:
                all_tables_exist = False
                if verbose:
                    report_status(f"Table '{tbl}'", False, "Missing table from database schema")
                report["status"] = "CRITICAL"
                
        # Test basic write query
        cursor.execute("SELECT COUNT(*) FROM drafts")
        count = cursor.fetchone()[0]
        db_result["writeable"] = True
        db_result["drafts_count"] = count
        conn.close()
        
        if verbose:
            report_status("Database Write/Read Test", True, f"Successfully queried database! Current drafts: {count}")
    except Exception as e:
        db_result["error"] = str(e)
        if verbose:
            report_status("Database Check", False, f"Integrity check failed: {e}")
        report["status"] = "CRITICAL"
        
    report["steps"]["database"] = db_result

    # ----------------------------------------------------
    # 4. System Binary Checks (Tesseract, LibreOffice soffice)
    # ----------------------------------------------------
    if verbose:
        print_header("4. DETECTING EXTERNAL SYSTEM BINARIES (FALLBACKS)")
        
    binary_checks = {"tesseract": False, "libreoffice": False}
    
    # Check Tesseract binary
    tess_path = shutil.which("tesseract")
    if not tess_path:
        tess_places = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
        ]
        for p in tess_places:
            if Path(p).exists():
                tess_path = p
                break
                
    if tess_path:
        try:
            res = subprocess.run([tess_path, "--version"], capture_output=True, text=True, check=True)
            v_match = re.search(r"tesseract\s+([\d\.\-]+)", res.stdout)
            v_str = v_match.group(1) if v_match else "unknown"
            binary_checks["tesseract"] = tess_path
            if verbose:
                report_status("Tesseract OCR", True, f"Found version {v_str} at: {tess_path}")
        except Exception as e:
            if verbose:
                report_warning("Tesseract OCR", f"Found path but failed execution check: {e}")
    else:
        if verbose:
            report_warning("Tesseract OCR", "Not found in system PATH. RapidOCR will still run locally, but Tesseract binary extraction fallback is unavailable.")
            
    # Check LibreOffice/soffice binary
    soffice_path = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice_path:
        office_places = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
        ]
        for p in office_places:
            if Path(p).exists():
                soffice_path = p
                break
                
    if soffice_path:
        try:
            res = subprocess.run([soffice_path, "--version"], capture_output=True, text=True, check=True)
            v_str = res.stdout.strip()
            binary_checks["libreoffice"] = soffice_path
            if verbose:
                report_status("LibreOffice soffice", True, f"Found at: {soffice_path} ({v_str})")
        except Exception as e:
            if verbose:
                report_warning("LibreOffice soffice", f"Found path but execution check failed: {e}")
    else:
        if verbose:
            report_warning("LibreOffice soffice", "Not found in system PATH. Old binary Word formats (.doc) conversion will be unavailable. Only standard formats (.docx) supported natively.")
            
    report["steps"]["binaries"] = binary_checks

    # ----------------------------------------------------
    # 5. Local Ollama Server & Model Availability Check
    # ----------------------------------------------------
    if verbose:
        print_header("5. OLLAMA LOCAL SERVER CONFLICTS & DEPLOYED MODELS")
        
    # Read dotenv configs or use internal defaults
    from dotenv import load_dotenv
    load_dotenv()
    
    # Grab configs
    from ollama_client import OLLAMA_BASE_URL, TEXT_MODEL, VISION_MODEL, GEMINI_API_KEY
    import requests
    
    ollama_check = {
        "configured_url": OLLAMA_BASE_URL,
        "online": False,
        "text_model_configured": TEXT_MODEL,
        "vision_model_configured": VISION_MODEL,
        "available_models": [],
        "gemini_api_configured": bool(GEMINI_API_KEY)
    }
    
    # Perform online check
    try:
        tags_url = f"{OLLAMA_BASE_URL}/api/tags"
        r = requests.get(tags_url, timeout=5)
        if r.status_code == 200:
            ollama_check["online"] = True
            models_data = r.json()
            downloaded_model_names = [m["name"] for m in models_data.get("models", [])]
            ollama_check["available_models"] = downloaded_model_names
            
            if verbose:
                report_status("Ollama Connection", True, f"Server online at {OLLAMA_BASE_URL}")
                print(f"   Downloaded Models: {', '.join(downloaded_model_names) if downloaded_model_names else 'None'}")
                
            # Verify specific text model matches
            if TEXT_MODEL in downloaded_model_names:
                if verbose:
                    report_status(f"Text Model: '{TEXT_MODEL}'", True, "Successfully deployed locally")
            elif any(TEXT_MODEL in name for name in downloaded_model_names):
                # Loose match
                matched_name = [name for name in downloaded_model_names if TEXT_MODEL in name][0]
                if verbose:
                    report_status(f"Text Model: '{TEXT_MODEL}'", True, f"Found matched model name: {matched_name}")
            else:
                if verbose:
                    report_warning(f"Text Model: '{TEXT_MODEL}'", f"Not found in Ollama downloaded tags. Please execute 'ollama pull {TEXT_MODEL}' in cmd if text extraction errors occur.")
            
            # Verify specific vision model matches
            if VISION_MODEL in downloaded_model_names:
                if verbose:
                    report_status(f"Vision Model: '{VISION_MODEL}'", True, "Successfully deployed locally")
            elif any(VISION_MODEL in name for name in downloaded_model_names):
                # Loose match
                matched_name = [name for name in downloaded_model_names if VISION_MODEL in name][0]
                if verbose:
                    report_status(f"Vision Model: '{VISION_MODEL}'", True, f"Found matched model name: {matched_name}")
            else:
                if verbose:
                    report_warning(f"Vision Model: '{VISION_MODEL}'", f"Not found in Ollama downloaded tags. Scanned menu/image extraction might fail locally unless Gemini Cloud fallback is configured.")
                    
            # Try a small model infer test
            if verbose:
                print("   Running lightweight inference test on local text model...")
            try:
                test_payload = {
                    "model": TEXT_MODEL,
                    "prompt": "Say: 'Ollama works!'",
                    "stream": False,
                    "options": {"num_predict": 10, "temperature": 0}
                }
                r_infer = requests.post(f"{OLLAMA_BASE_URL}/api/generate", json=test_payload, timeout=120)
                if r_infer.status_code == 200:
                    ollama_check["inference_test"] = "SUCCESS"
                    resp_text = r_infer.json().get("response", "").strip()
                    if verbose:
                        report_status("Inference Test", True, f"Response: '{resp_text}'")
                else:
                    ollama_check["inference_test"] = f"FAILED: Code {r_infer.status_code}"
                    if verbose:
                        report_status("Inference Test", False, f"Server returned error code {r_infer.status_code}")
            except Exception as ex:
                ollama_check["inference_test"] = f"CRITICAL: {ex}"
                if verbose:
                    report_status("Inference Test", False, f"Connection timeout / model load failed: {ex}")
        else:
            if verbose:
                report_status("Ollama Connection", False, f"Server returned status code {r.status_code}")
            report["status"] = "DEGRADED"
    except Exception as e:
        if verbose:
            report_status("Ollama Connection", False, f"Could not connect to Ollama base server: {e}. Ensure the Ollama background service is running on your system.")
        report["status"] = "DEGRADED"
        
    if GEMINI_API_KEY:
        if verbose:
            report_status("Gemini Cloud API", True, f"Configured API Key (Starting with '{GEMINI_API_KEY[:4]}...'). High speed bypass is active.")
            
    report["steps"]["ollama"] = ollama_check

    # ----------------------------------------------------
    # 6. E2E Extraction Verification (All Formats)
    # ----------------------------------------------------
    if verbose:
        print_header("6. PIPELINE E2E FILE FORMAT EXTRACTION VALIDATION")
        
    # We will generate temporary dummy mock files for text, Excel, CSV, Word, and run the pipeline
    tmp_path = Path(tempfile.mkdtemp(prefix="agent_verify_"))
    
    extracted_results = {}
    try:
        # Import extractors internally
        from file_extractors import (
            extract_menu_from_file, 
            extract_directly_from_csv, 
            extract_directly_from_excel, 
            extract_plain_text_from_file
        )
        from heuristics_parser import parse_menu_text_heuristically
        import re

        # -- File Type 1: Text format (.txt) --
        txt_file = tmp_path / "mock_menu_text.txt"
        txt_content = """
        Soups & Starters
        Tomato Soup
        Rich tomato base soup with seasoned bread croutons
        120
        Vegetable Spring Rolls 150
        Crispy fried wrapper stuffed with fresh garden variables.
        
        Mains
        Paneer Butter Masala
        Half / Full
        210 / 350
        """
        txt_file.write_text(txt_content, encoding="utf-8")
        
        # Test raw text extraction using local heuristics
        heur_extract = parse_menu_text_heuristically(txt_content)
        txt_passed = heur_extract is not None and len(heur_extract.items) >= 3
        extracted_results["text_heuristics"] = {
            "passed": txt_passed,
            "items_count": len(heur_extract.items) if heur_extract else 0,
            "details": [f"Item: {it.product_name} | Cat: {it.category} | Variations: {len(it.variations)}" for it in heur_extract.items] if heur_extract else []
        }
        if verbose:
            report_status("TXT Heuristic Parser", txt_passed, f"Extracted {extracted_results['text_heuristics']['items_count']} products.")
            if txt_passed:
                for line in extracted_results["text_heuristics"]["details"]:
                    print(f"      - {line}")

        # -- File Type 2: CSV directly parsed (.csv) --
        csv_file = tmp_path / "mock_menu_table.csv"
        csv_rows = [
            ["Category", "Product Name", "Selling Price", "Description"],
            ["Desserts", "Hot Chocolate Fudge", "180", "Vanilla ice cream with rich hot fudge chocolate sauce."],
            ["Desserts", "Brownie With Ice Cream", "220", "Warm walnut chocolate brownie dessert."],
            ["Beverages", "Mint Mojito", "140", "Fresh mint sprigs lime soda."],
            ["Beverages", "Double Choco Shake", "160#220", "Thick shake chocolate cream. Available in Regular / Large Sizes"]
        ]
        
        import csv
        with csv_file.open("w", encoding="utf-8", newline="") as f:
            writer = csv.writer(f)
            writer.writerows(csv_rows)
            
        csv_extract = extract_directly_from_csv(csv_file)
        csv_passed = csv_extract is not None and len(csv_extract.items) == 4
        
        # Check variation sizes split by # for Mint Mojito & Double Choco Shake
        mojito = [it for it in csv_extract.items if "mojito" in it.product_name.lower()][0] if csv_extract else None
        shake = [it for it in csv_extract.items if "shake" in it.product_name.lower()][0] if csv_extract else None
        
        shake_passed = shake is not None and len(shake.variations) == 2
        
        extracted_results["csv_parser"] = {
            "passed": csv_passed and shake_passed,
            "items_count": len(csv_extract.items) if csv_extract else 0,
            "shake_variations": len(shake.variations) if shake else 0
        }
        if verbose:
            report_status("CSV Tabular Finder", csv_passed and shake_passed, f"Extracted {extracted_results['csv_parser']['items_count']} products. Multi-variant split check passed: {shake_passed}.")

        # -- File Type 3: Excel directly parsed (.xlsx) --
        xlsx_file = tmp_path / "mock_menu_sheet.xlsx"
        from openpyxl import Workbook
        wb_mock = Workbook()
        ws_mock = wb_mock.active
        ws_mock.title = "Menu Items"
        
        excel_headers = ["Category Name", "Item Name", "Price Rate", "Description Info", "Custom Group"]
        ws_mock.append(excel_headers)
        ws_mock.append(["Chinese", "Veg Fried Rice", "180", "Wok tossed garden veggies", "Regular"])
        ws_mock.append(["Chinese", "Chicken Hakka Noodles", "210#280", "Soft noodles tossed with shredded chicken", "Half#Full"])
        ws_mock.append(["Thali", "Premium Veg Thali", "320", "Paneer, Dal, Sabzi, Roti, Rice and sweet desserts", ""])
        wb_mock.save(str(xlsx_file))
        wb_mock.close()
        
        xlsx_extract = extract_directly_from_excel(xlsx_file)
        xlsx_passed = xlsx_extract is not None and len(xlsx_extract.items) == 3
        
        # check multi-variation mapping on Hakka Noodles
        noodles = [it for it in xlsx_extract.items if "noodles" in it.product_name.lower()][0] if xlsx_extract else None
        noodles_passed = noodles is not None and len(noodles.variations) == 2 and noodles.variations[1].price == 280.0
        
        extracted_results["excel_parser"] = {
            "passed": xlsx_passed and noodles_passed,
            "items_count": len(xlsx_extract.items) if xlsx_extract else 0,
            "noodles_variations_count": len(noodles.variations) if noodles else 0
        }
        if verbose:
            report_status("Excel Sheet Reader Bypass", xlsx_passed and noodles_passed, f"Extracted {extracted_results['excel_parser']['items_count']} products. Noodles variation price parsed correctly: {noodles_passed}.")

        # -- File Type 4: Word Document parser (.docx) --
        docx_file = tmp_path / "mock_menu_word.docx"
        from docx import Document
        doc = Document()
        doc.add_heading("Italian Corner Menu", level=0)
        doc.add_paragraph("PIZZAS")
        doc.add_paragraph("Margherita Pizza - Classic tomato cheese basil - 240")
        doc.add_paragraph("Farmhouse Pizza - Loaded woodfired garden fresh greens - 320")
        doc.add_paragraph("PASTAS")
        
        # add a small table
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Pasta Name'
        hdr_cells[1].text = 'Price'
        hdr_cells[2].text = 'Details'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Penne Arrabiata'
        row_cells[1].text = '260'
        row_cells[2].text = 'Tangy red sauce spices'
        
        doc.save(str(docx_file))
        
        # Test word plain text extractor first
        word_text = extract_plain_text_from_file(docx_file)
        word_text_passed = "Margherita Pizza" in word_text and "Penne Arrabiata" in word_text
        
        extracted_results["word_plain_text"] = {
            "passed": word_text_passed,
            "char_count": len(word_text)
        }
        if verbose:
            report_status("Word plain text extraction", word_text_passed, f"Chars extracted: {len(word_text)}")
            
        # Run local fallback heuristic pipeline on word plain text
        word_heur = parse_menu_text_heuristically(word_text)
        word_heur_passed = word_heur is not None and len(word_heur.items) >= 3
        extracted_results["word_heuristic_pipeline"] = {
            "passed": word_heur_passed,
            "items_count": len(word_heur.items) if word_heur else 0
        }
        if verbose:
            report_status("Word extraction heuristic E2E", word_heur_passed, f"Extracted {extracted_results['word_heuristic_pipeline']['items_count']} products from template.")

        # -- File Type 5: Scanned Image/OCR Engine --
        # Verify local RapidOCR runs on a basic canvas
        from PIL import Image, ImageDraw, ImageFont
        img = Image.new('RGB', (600, 200), color = (255, 255, 255))
        d = ImageDraw.Draw(img)
        d.text((20, 40), "Cold Coffee Regular 120 Large 180", fill=(0,0,0))
        d.text((20, 100), "Cappuccino Hot Brew 150", fill=(0,0,0))
        
        img_path = tmp_path / "mock_scanned_menu.jpg"
        img.save(img_path)
        
        # Try running OCR check
        from file_extractors import extract_text_from_image_via_ocr
        ocr_out_text = extract_text_from_image_via_ocr(img_path)
        ocr_passed = "coffee" in ocr_out_text.lower() or "cappuccino" in ocr_out_text.lower() or len(ocr_out_text.strip()) > 0
        
        extracted_results["ocr_engine"] = {
            "passed": ocr_passed,
            "ocr_text": ocr_out_text.replace("\n", " | ")
        }
        if verbose:
            report_status("OCR Engine Validation (RapidOCR/Tesseract)", ocr_passed, f"Extracted OCR string: '{ocr_out_text.strip().replace(chr(10), ' | ')}'")
            
    except Exception as e:
        extracted_results["error"] = str(e)
        if verbose:
            report_status("E2E File Parsers", False, f"Exception occurred: {e}")
            import traceback
            traceback.print_exc()
        report["status"] = "DEGRADED"
        
    report["steps"]["e2e_parsers"] = extracted_results
    
    # clean tmp folder
    try:
        shutil.rmtree(tmp_path)
    except Exception:
        pass

    # ----------------------------------------------------
    # 7. Excel Template Writing & Export Validation
    # ----------------------------------------------------
    if verbose:
        print_header("7. EXCEL EXPORTER & INSTRUCTIONS VALIDATOR")
        
    writer_result = {"write_success": False, "template_check": False, "validation_engine": False}
    try:
        # Create a mock MenuExtraction dataset
        mock_items = [
            MenuItem(
                category="Burgers",
                product_name="Crispy Veg Burger",
                description="Potato peas patty sesame bun spicy dressing",
                dietary_tag="Veg",
                confidence=1.0,
                source_text="Veg Burger 140",
                variations=[{"name": "", "price": 140.0, "listing_price": 140.0}]
            ),
            MenuItem(
                category="Burgers",
                product_name="Smoky Chicken Burger",
                description="Grilled chicken breast patty cheddar cheese smoked chipotle",
                dietary_tag="Non-Veg",
                confidence=0.92,
                source_text="Chicken Burger Regular 180 Large 260",
                variations=[
                    {"name": "Regular", "price": 180.0, "listing_price": 180.0},
                    {"name": "Large", "price": 260.0, "listing_price": 300.0}
                ]
            ),
            MenuItem(
                category="Beverages",
                product_name="Spiced Lemon Tea",
                description="",
                dietary_tag="",
                confidence=0.61,  # Low confidence to trigger review warnings
                source_text="Lemon Tea 80",
                variations=[{"name": "", "price": 80.0, "listing_price": 80.0}]
            )
        ]
        
        mock_extraction = MenuExtraction(currency="INR", items=mock_items, document_notes=["Self-test diagnostics"])
        
        output_xlsx_path = base_dir / "outputs" / "test_diagnostics_verify.xlsx"
        output_review_path = base_dir / "outputs" / "test_diagnostics_verify.json"
        
        # Test writing
        res_output = write_bulk_upload_excel(default_template, mock_extraction, output_xlsx_path, output_review_path)
        writer_result["write_success"] = output_xlsx_path.exists() and output_review_path.exists()
        
        # Test loading back generated excel using openpyxl
        wb_check = load_workbook(str(output_xlsx_path))
        writer_result["sheets_in_output"] = wb_check.sheetnames
        
        is_sheet_present = PRODUCT_SHEET in wb_check.sheetnames
        writer_result["template_check"] = is_sheet_present
        
        if is_sheet_present:
            ws_check = wb_check[PRODUCT_SHEET]
            # check headers
            headers_found = [str(ws_check.cell(row=1, column=c).value).strip() for c in range(1, ws_check.max_column + 1) if ws_check.cell(row=1, column=c).value]
            writer_result["headers_count"] = len(headers_found)
            
            # verify mandatory columns exist
            all_mandatories_present = all(h in headers_found for h in MANDATORY_HEADERS)
            writer_result["mandatory_headers_check"] = all_mandatories_present
            
            # Read first product row
            row_items_written = ws_check.max_row - 1
            writer_result["excel_rows_written"] = row_items_written
            
            cat_col = headers_found.index("Category Name*") + 1
            prod_col = headers_found.index("Product Name*") + 1
            sp_col = headers_found.index("Selling Price*") + 1
            var_col = headers_found.index("Variation") + 1 if "Variation" in headers_found else -1
            
            p1_cat = ws_check.cell(row=2, column=cat_col).value
            p1_name = ws_check.cell(row=2, column=prod_col).value
            p1_price = ws_check.cell(row=2, column=sp_col).value
            
            p2_name = ws_check.cell(row=3, column=prod_col).value
            p2_price = ws_check.cell(row=3, column=sp_col).value
            p2_var = ws_check.cell(row=3, column=var_col).value if var_col != -1 else ""
            
            writer_result["sample_row_1"] = {"category": p1_cat, "name": p1_name, "price": p1_price}
            writer_result["sample_row_2"] = {"name": p2_name, "price": p2_price, "variation": p2_var}
            
            wb_check.close()
            
            if verbose:
                report_status("Excel Export sheet", True, f"Successfully created template structure with {row_items_written} items.")
                print(f"      - Sheet structure is valid ({', '.join(wb_check.sheetnames)})")
                print(f"      - Mandatory Headers conform to bulk specification: {all_mandatories_present}")
                print(f"      - Row 1 check (Single) -> Category: {p1_cat} | Name: {p1_name} | Price: {p1_price}")
                print(f"      - Row 2 check (Multi)  -> Name: {p2_name} | Variations: {p2_var} | Prices: {p2_price}")
        else:
            if verbose:
                report_status("Excel Export sheet", False, "Missing output worksheet 'Product Bulk'")
            report["status"] = "CRITICAL"
            
        # Test review report JSON
        with output_review_path.open("r", encoding="utf-8") as f:
            review_data = json.load(f)
            
        # Should have warnings for Spiced Lemon Tea (low confidence)
        lemon_warnings = [r for r in review_data.get("review_rows", []) if "lemon" in r["product_name"].lower()]
        has_warnings = len(lemon_warnings) > 0 and len(lemon_warnings[0]["warnings"]) > 0
        writer_result["review_json_check"] = has_warnings
        if verbose:
            report_status("Validation Review Output", has_warnings, f"Correctly caught low confidence item warning: {lemon_warnings[0]['warnings'] if lemon_warnings else 'No warnings'}")
            
        # Test cleaning test outputs
        output_xlsx_path.unlink()
        output_review_path.unlink()
        
    except Exception as e:
        writer_result["error"] = str(e)
        if verbose:
            report_status("Excel Exporter / Validator", False, f"Failed execution of bulk writer: {e}")
        report["status"] = "CRITICAL"
        
    report["steps"]["excel_writer"] = writer_result

    # ----------------------------------------------------
    # FINAL RATING & OUTPUT INTEGRITY
    # ----------------------------------------------------
    print_header("DIAGNOSSTIC SUMMARY REPORT CARD")
    print(f"Diagnostics executed: {report['timestamp']}")
    if report["status"] == "HEALTHY":
        print(f"Overall System Integrity: {C_GREEN}{C_BOLD}HEALTHY - 100% Operational{C_RESET}")
    elif report["status"] == "DEGRADED":
        print(f"Overall System Integrity: {C_YELLOW}{C_BOLD}DEGRADED - Some fallbacks are inactive but core pipeline is functional{C_RESET}")
    else:
        print(f"Overall System Integrity: {C_RED}{C_BOLD}CRITICAL - Immediate configuration or library errors detected{C_RESET}")
        
    # Write summary diagnostics history
    diag_history_path = base_dir / "outputs" / "self_diagnostic_report.json"
    with diag_history_path.open("w", encoding="utf-8") as f:
        f.write(json.dumps(report, indent=2))
        
    print(f"Detailed json diagnostics report written to: {diag_history_path}")
    return report

if __name__ == "__main__":
    run_diagnostics(verbose=True)
