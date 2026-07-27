from __future__ import annotations



import csv
import concurrent.futures
import shutil
import subprocess
import tempfile
import re
from pathlib import Path
from typing import List, Optional

from PIL import Image

import sys
from pathlib import Path
backend_dir = Path(__file__).resolve().parent
if str(backend_dir) not in sys.path:
    sys.path.insert(0, str(backend_dir))

import os
from ollama_client import (
    extract_from_image_with_ollama, 
    extract_from_text_with_ollama, 
    merge_extractions,
    extract_from_text_with_gemini,
    extract_from_images_with_gemini
)
from heuristics_parser import parse_menu_text_heuristically
from schemas import MenuExtraction, MenuItem

SUPPORTED_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}
SUPPORTED_PDF_EXTS = {".pdf"}
SUPPORTED_EXCEL_EXTS = {".xlsx", ".xlsm", ".xls", ".csv"}
SUPPORTED_WORD_EXTS = {".docx", ".doc"}
SUPPORTED_TEXT_EXTS = {".txt"}
SUPPORTED_EXTS = SUPPORTED_IMAGE_EXTS | SUPPORTED_PDF_EXTS | SUPPORTED_EXCEL_EXTS | SUPPORTED_WORD_EXTS | SUPPORTED_TEXT_EXTS


def _read_csv(path: Path) -> str:
    lines = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            cells = [str(c).strip() for c in row if str(c).strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _convert_doc_to_docx_with_libreoffice(path: Path) -> Path:
    """
    Optional old .doc support.
    Works only if LibreOffice/soffice is installed and available in PATH.
    For best reliability, ask clients to upload .docx or PDF.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise ValueError(
            "Old .doc Word files need LibreOffice installed for conversion. "
            "Please upload .docx or PDF, or install LibreOffice and ensure soffice is in PATH."
        )

    out_dir = Path(tempfile.mkdtemp(prefix="menu_doc_convert_"))
    subprocess.run(
        [soffice, "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(path)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    converted = out_dir / f"{path.stem}.docx"
    if not converted.exists():
        matches = list(out_dir.glob("*.docx"))
        if not matches:
            raise ValueError("Could not convert .doc file to .docx.")
        converted = matches[0]
    return converted


def _read_excel(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix == ".csv":
        return _read_csv(path)

    # .xlsx/.xlsm is handled by openpyxl. Old .xls is handled by pandas+xlrd.
    if suffix in {".xlsx", ".xlsm"}:
        from openpyxl import load_workbook

        wb = load_workbook(str(path), read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f"SHEET: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)

    if suffix == ".xls":
        import pandas as pd

        workbook = pd.read_excel(str(path), sheet_name=None, header=None, dtype=str, engine="xlrd")
        lines = []
        for sheet_name, df in workbook.items():
            lines.append(f"SHEET: {sheet_name}")
            df = df.fillna("")
            for _, row in df.iterrows():
                cells = [str(c).strip() for c in row.tolist() if str(c).strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)

    return ""


def extract_plain_text_from_file(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")

    if suffix in SUPPORTED_EXCEL_EXTS:
        return _read_excel(path)

    if suffix == ".docx":
        return _read_docx(path)

    if suffix == ".doc":
        return _read_docx(_convert_doc_to_docx_with_libreoffice(path))

    if suffix == ".pdf":
        import fitz  # PyMuPDF
        import re

        doc = fitz.open(str(path))
        parts = []

        category_keywords = {
            "menu", "drinks", "beverages", "desserts", "starters", "mains", "appetizers", 
            "sides", "soups", "salads", "breads", "curry", "signature", "breakfast", 
            "hi tea", "base", "protein", "chutney", "pickle", "seeds", "add on", "bases", "curries",
            "quick bites", "momos", "burgers", "fries", "shake", "shakes", "pizza", "noodles", "fried rice"
        }

        price_line_regex = re.compile(
            r'^\s*(?:\$|Rs\.?|₹|INR|\+)?\s*(\d+(?:\.\d{1,2})?)\s*(?:[\/|\\\,]\s*(?:\$|Rs\.?|₹|INR|\+)?\s*(\d+(?:\.\d{1,2})?))*\s*(?:\+?\s*(?:\$|Rs\.?|₹|INR)?\s*\d+(?:\.\d{1,2})?)?\s*$',
            re.IGNORECASE
        )

        def is_price_block(block_text: str) -> bool:
            lines = [l.strip() for l in block_text.splitlines() if l.strip()]
            if not lines:
                return False
            price_lines = 0
            for l in lines:
                cleaned = re.sub(r'^[-\+\.\s\:\,\;\*\/\|]+', '', l)
                cleaned = re.sub(r'[-\+\.\s\:\,\;\*\/\|]+$', '', cleaned).strip()
                if price_line_regex.match(cleaned) or re.search(r'^\d+$', cleaned):
                    price_lines += 1
            return price_lines >= 0.7 * len(lines)

        for page_num, page in enumerate(doc, start=1):
            blocks = page.get_text("blocks")
            if not blocks:
                continue

            text_blocks = []
            for b in blocks:
                if b[6] == 0:
                    text = b[4].strip()
                    if text:
                        text_blocks.append({
                            "x0": b[0], "y0": b[1], "x1": b[2], "y1": b[3],
                            "text": b[4], "lines": [l.strip() for l in b[4].splitlines() if l.strip()],
                            "merged": False
                        })

            num_blocks = len(text_blocks)
            page_width = page.rect.width

            for i in range(num_blocks):
                if text_blocks[i]["merged"]:
                    continue
                b1 = text_blocks[i]
                if is_price_block(b1["text"]):
                    continue

                price_candidates = []
                for j in range(num_blocks):
                    if i == j or text_blocks[j]["merged"]:
                        continue
                    b2 = text_blocks[j]

                    if b2["x0"] > b1["x0"]:
                        h_gap = b2["x0"] - b1["x1"]
                        y_overlap = min(b1["y1"], b2["y1"]) - max(b1["y0"], b2["y0"])
                        b1_h = b1["y1"] - b1["y0"]
                        b2_h = b2["y1"] - b2["y0"]
                        min_h = min(b1_h, b2_h)

                        if min_h > 0:
                            overlap_ratio = y_overlap / min_h
                        else:
                            overlap_ratio = 0

                        if h_gap < 130 and overlap_ratio > 0.45:
                            if is_price_block(b2["text"]):
                                price_candidates.append((j, b2))

                price_candidates = sorted(price_candidates, key=lambda pair: pair[1]["y0"])

                for p_idx, b2 in price_candidates:
                    lines1 = b1["lines"]
                    lines2 = b2["lines"]

                    item_indices = []
                    for idx, l in enumerate(lines1):
                        cleaned = re.sub(r'^[-\+\.\s\:\,\;\*\/\|]+', '', l)
                        cleaned = re.sub(r'[-\+\.\s\:\,\;\*\/\|]+$', '', cleaned).strip()
                        is_hdr = False
                        if len(cleaned) < 40:
                            if cleaned.isupper() and not price_line_regex.match(cleaned):
                                is_hdr = True
                            elif cleaned.lower() in category_keywords:
                                is_hdr = True
                        if re.search(r'\d+$', cleaned):
                            is_hdr = True
                        if not is_hdr:
                            item_indices.append(idx)

                    merged_map = {}
                    for count, idx in enumerate(item_indices):
                        l1 = lines1[idx]
                        l2 = lines2[count] if count < len(lines2) else ""
                        if l2:
                            merged_map[idx] = f"{l1} {l2}"

                    new_lines = []
                    for idx in range(len(lines1)):
                        if idx in merged_map:
                            new_lines.append(merged_map[idx])
                        else:
                            new_lines.append(lines1[idx])

                    if len(lines2) > len(item_indices):
                        for extra_idx in range(len(item_indices), len(lines2)):
                            new_lines.append(lines2[extra_idx])

                    b1["lines"] = new_lines
                    b1["text"] = "\n".join(new_lines)
                    b1["x1"] = max(b1["x1"], b2["x1"])
                    b1["y0"] = min(b1["y0"], b2["y0"])
                    b1["y1"] = max(b1["y1"], b2["y1"])
                    b2["merged"] = True

            active_blocks = [b for b in text_blocks if not b["merged"]]
            if not active_blocks:
                continue

            narrow_blocks = [b for b in active_blocks if (b["x1"] - b["x0"]) < 0.65 * page_width]
            wide_blocks = [b for b in active_blocks if (b["x1"] - b["x0"]) >= 0.65 * page_width]

            sorted_narrow = sorted(narrow_blocks, key=lambda b: b["x0"])
            columns = []
            for b in sorted_narrow:
                b_x0 = b["x0"]
                placed = False
                for col in columns:
                    col_x0_avg = sum(cb["x0"] for cb in col) / len(col)
                    if abs(b_x0 - col_x0_avg) < page_width * 0.12:
                        col.append(b)
                        placed = True
                        break
                if not placed:
                    columns.append([b])

            columns = sorted(columns, key=lambda col: sum(cb["x0"] for cb in col) / len(col))
            page_lines = []

            for c_idx, col in enumerate(columns):
                col_sorted = sorted(col, key=lambda b: b["y0"])
                for b in col_sorted:
                    page_lines.extend(b["lines"])

            if wide_blocks:
                wide_sorted = sorted(wide_blocks, key=lambda b: b["y0"])
                for b in wide_sorted:
                    page_lines.extend(b["lines"])

            text = "\n".join(page_lines).strip()
            if text:
                parts.append(f"--- PAGE {page_num} ---\n{text}")
        return "\n".join(parts)

    return ""


def _compress_image_for_vision(input_path: Path, output_path: Path, max_size: int = 1800) -> Path:
    img = Image.open(input_path).convert("RGB")
    img.thumbnail((max_size, max_size))
    img.save(output_path, format="JPEG", quality=92)
    return output_path


def _pdf_pages_to_images(path: Path, max_pages: int = 20) -> List[Path]:
    import fitz  # PyMuPDF

    out_paths: List[Path] = []
    temp_dir = Path(tempfile.mkdtemp(prefix="menu_pdf_pages_"))
    doc = fitz.open(str(path))
    for index, page in enumerate(doc):
        if index >= max_pages:
            break
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        out_path = temp_dir / f"page_{index + 1}.jpg"
        pix.save(str(out_path))
        out_paths.append(out_path)
    return out_paths


def _chunk_text_by_lines(text: str, max_chars: int = 4000) -> List[str]:
    chunks = []
    current_chunk = []
    current_len = 0
    for line in text.splitlines():
        line_len = len(line) + 1
        if current_len + line_len > max_chars and current_chunk:
            chunks.append("\n".join(current_chunk))
            current_chunk = []
            current_len = 0
        current_chunk.append(line)
        current_len += line_len
    if current_chunk:
        chunks.append("\n".join(current_chunk))
    return chunks


def extract_directly_from_excel(path: Path) -> Optional[MenuExtraction]:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(path), data_only=True)
        
        # 1. Identify sheets
        items_sheet_name = None
        modifiers_sheet_name = None
        
        for sname in wb.sheetnames:
            lower_name = sname.lower()
            if "modifier" in lower_name or "addon" in lower_name or "customization" in lower_name:
                modifiers_sheet_name = sname
            elif "item" in lower_name or "menu" in lower_name or "product" in lower_name:
                if not items_sheet_name:
                    items_sheet_name = sname
                    
        if not items_sheet_name:
            items_sheet_name = wb.sheetnames[0]
            
        # 2. Parse Items sheet
        ws_items = wb[items_sheet_name]
        rows = list(ws_items.iter_rows(values_only=True))
        if not rows:
            return None
            
        header_row_idx = -1
        col_map = {}
        best_score = 0
        
        for idx, r in enumerate(rows[:20]):
            row_score = 0
            temp_map = {}
            for c_idx, val in enumerate(r):
                if val is None:
                    continue
                val_str = str(val).strip().lower()
                
                if ("category" in val_str or "section" in val_str or "group" in val_str) and "category" not in temp_map:
                    if "add on" not in val_str:
                        temp_map["category"] = c_idx
                        row_score += 1
                if ("name" in val_str or "product" in val_str or "item" in val_str) and "category" not in val_str and "product_name" not in temp_map:
                    if "add on" not in val_str:
                        temp_map["product_name"] = c_idx
                        row_score += 2
                if ("price" in val_str or "rate" in val_str or "mrp" in val_str or "cost" in val_str) and "price" not in temp_map:
                    if "cost" not in val_str:
                        temp_map["price"] = c_idx
                        row_score += 2
                if ("desc" in val_str or "detail" in val_str or "info" in val_str) and "description" not in temp_map:
                    temp_map["description"] = c_idx
                    row_score += 1
                if "variation" in val_str and "group" not in val_str and "variation" not in temp_map:
                    temp_map["variation"] = c_idx
                    row_score += 1
                if ("variant group" in val_str or "customization group" in val_str) and "variant_group" not in temp_map:
                    temp_map["variant_group"] = c_idx
                    row_score += 1
                    
            if "product_name" in temp_map and "price" in temp_map and row_score > best_score:
                best_score = row_score
                header_row_idx = idx
                col_map = temp_map
                
        if header_row_idx == -1:
            category_idx = 0
            name_idx = 1
            price_idx = 2
            desc_idx = 3
            var_idx = -1
            group_idx = -1
            header_row_idx = 0
        else:
            category_idx = col_map.get("category", -1)
            name_idx = col_map["product_name"]
            price_idx = col_map["price"]
            desc_idx = col_map.get("description", -1)
            var_idx = col_map.get("variation", -1)
            group_idx = col_map.get("variant_group", -1)
            
        items_list = []
        items_dict = {} # Key: (category, product_name), Value: item dict
        
        current_category = "Uncategorized"
        for r in rows[header_row_idx + 1:]:
            if not any(c is not None for c in r):
                continue
                
            name_val = r[name_idx] if name_idx < len(r) else None
            if name_val is None or str(name_val).strip() == "" or str(name_val).strip().lower() in ["item name", "product name", "name", "none"]:
                continue
                
            if category_idx != -1 and category_idx < len(r) and r[category_idx] is not None:
                cat_val = str(r[category_idx]).strip()
                if cat_val:
                    current_category = cat_val
                    
            price_val = r[price_idx] if price_idx < len(r) else None
            var_val = str(r[var_idx]).strip() if var_idx != -1 and var_idx < len(r) and r[var_idx] is not None else ""
            group_val = str(r[group_idx]).strip() if group_idx != -1 and group_idx < len(r) and r[group_idx] is not None else ""
            
            price_str = str(price_val).strip() if price_val is not None else ""
            
            # Check cell for inline # names
            hash_cell_names = []
            for cell_val in r:
                if cell_val is not None and "#" in str(cell_val):
                    cell_str = str(cell_val).strip()
                    if cell_str != price_str and not re.search(r'^\d*(?:\.\d+)?(?:#\d*(?:\.\d+)?)*$', cell_str):
                        hash_cell_names = [x.strip() for x in cell_str.split("#") if x.strip()]
                        break
            
            row_variations = []
            if "#" in price_str:
                prices_split = [x.strip() for x in price_str.split("#") if x.strip()]
                if hash_cell_names and len(hash_cell_names) == len(prices_split):
                    v_names = hash_cell_names
                else:
                    if len(prices_split) == 2:
                        v_names = ["Half", "Full"]
                    elif len(prices_split) == 3:
                        v_names = ["Small", "Medium", "Large"]
                    else:
                        v_names = [f"Size {k+1}" for k in range(len(prices_split))]
                
                for k, p_str in enumerate(prices_split):
                    try:
                        p_val = float(p_str)
                    except ValueError:
                        p_val = 0.0
                    row_variations.append({"name": v_names[k], "price": p_val, "listing_price": p_val})
            else:
                try:
                    price_num = float(price_val) if price_val is not None and price_str != "" else 0.0
                except ValueError:
                    price_num = 0.0
                row_variations.append({"name": var_val, "price": price_num, "listing_price": price_num})
                
            desc_val = str(r[desc_idx]).strip() if desc_idx != -1 and desc_idx < len(r) and r[desc_idx] is not None else ""
            item_name = str(name_val).strip()
            
            key = (current_category.lower().strip(), item_name.lower().strip())
            
            if key in items_dict:
                existing_item = items_dict[key]
                for new_v in row_variations:
                    exists = False
                    for ex_v in existing_item["variations"]:
                        if ex_v["name"].lower() == new_v["name"].lower():
                            # Update price if it matches but might have been imported with fallback previously
                            if ex_v["price"] == 0.0 and new_v["price"] != 0.0:
                                ex_v["price"] = new_v["price"]
                                ex_v["listing_price"] = new_v["listing_price"]
                            exists = True
                            break
                    if not exists:
                        existing_item["variations"].append(new_v)
                if group_val and not existing_item.get("variant_group_name"):
                    existing_item["variant_group_name"] = group_val
            else:
                item_dict = {
                    "category": current_category,
                    "product_name": item_name,
                    "description": desc_val,
                    "variant_group_name": group_val,
                    "variations": row_variations,
                    "modifiers_groups": {}
                }
                items_list.append(item_dict)
                items_dict[key] = item_dict
                
        # 3. Parse Modifiers sheet if exists
        if modifiers_sheet_name:
            ws_mods = wb[modifiers_sheet_name]
            mod_rows = list(ws_mods.iter_rows(values_only=True))
            if mod_rows:
                m_col_map = {}
                best_mod_score = 0
                for r_idx, r in enumerate(mod_rows[:20]):
                    temp_map = {}
                    row_score = 0
                    for c_idx, val in enumerate(r):
                        if val is None:
                            continue
                        val_str = str(val).strip().lower()
                        if ("item" in val_str or "product" in val_str or "parent" in val_str) and "parent_name" not in temp_map:
                            temp_map["parent_name"] = c_idx
                            row_score += 2
                        if ("group" in val_str or "customization" in val_str) and "group_name" not in temp_map:
                            temp_map["group_name"] = c_idx
                            row_score += 1
                        if ("modifier" in val_str or "option" in val_str or "addon name" in val_str) and "modifier_name" not in temp_map:
                            temp_map["modifier_name"] = c_idx
                            row_score += 2
                        if ("price" in val_str or "addon" in val_str) and "price_addon" not in temp_map:
                            temp_map["price_addon"] = c_idx
                            row_score += 2
                    
                    if "parent_name" in temp_map and "modifier_name" in temp_map and row_score > best_mod_score:
                        best_mod_score = row_score
                        m_col_map = temp_map
                
                p_name_idx = m_col_map.get("parent_name", 1)
                group_idx = m_col_map.get("group_name", 2)
                mod_idx = m_col_map.get("modifier_name", 3)
                addon_idx = m_col_map.get("price_addon", 4)
                
                header_row_mod = 0
                for idx, r in enumerate(mod_rows[:20]):
                    if r[p_name_idx] and "item" in str(r[p_name_idx]).lower():
                        header_row_mod = idx
                        break
                        
                for m_r in mod_rows[header_row_mod + 1:]:
                    if not any(c is not None for c in m_r):
                        continue
                    p_name_val = m_r[p_name_idx] if p_name_idx < len(m_r) else None
                    if not p_name_val:
                        continue
                        
                    p_name_clean = str(p_name_val).lower().strip()
                    # Match by name in items list (since category might not match or be absent in modifier sheet)
                    matched_items = [it for it in items_list if it["product_name"].lower().strip() == p_name_clean]
                    for item_dict in matched_items:
                        g_name = str(m_r[group_idx]).strip() if group_idx < len(m_r) and m_r[group_idx] is not None else "Options"
                        m_name = str(m_r[mod_idx]).strip() if mod_idx < len(m_r) and m_r[mod_idx] is not None else ""
                        m_addon = m_r[addon_idx] if addon_idx < len(m_r) else 0.0
                        try:
                            m_addon_val = float(m_addon) if m_addon is not None and str(m_addon).strip() != "" else 0.0
                        except ValueError:
                            m_addon_val = 0.0
                            
                        if g_name not in item_dict["modifiers_groups"]:
                            item_dict["modifiers_groups"][g_name] = []
                        item_dict["modifiers_groups"][g_name].append({
                            "name": m_name,
                            "addon": m_addon_val
                        })
                        
            # Apply modifiers to variations
            for item in items_list:
                if item["modifiers_groups"]:
                    selected_group = None
                    groups_list = list(item["modifiers_groups"].keys())
                    
                    for g in groups_list:
                        gl = g.lower()
                        if "portion" in gl or "size" in gl or "protein" in gl or "option" in gl:
                            selected_group = g
                            break
                    if not selected_group:
                        selected_group = groups_list[0]
                        
                    mods = item["modifiers_groups"][selected_group]
                    base_price = item["variations"][0]["price"] if item["variations"] else 0.0
                    
                    new_variations = []
                    for m in mods:
                        var_price = base_price + m["addon"]
                        new_variations.append({
                            "name": m["name"],
                            "price": var_price,
                            "listing_price": var_price
                        })
                    item["variations"] = new_variations
                    
        wb.close()
        
        # Build MenuItem validation payload
        all_items = []
        for item in items_list:
            vars_list = item["variations"]
            if len(vars_list) > 1:
                non_empty_vars = [v for v in vars_list if v["name"] != ""]
                if non_empty_vars:
                    vars_list = non_empty_vars
                    
            item_dict = {
                "category": item["category"],
                "product_name": item["product_name"],
                "description": item["description"],
                "dietary_tag": "",
                "confidence": 1.0,
                "source_text": f"Direct spreadsheet import",
                "variations": [
                    {
                        "name": v["name"],
                        "price": v["price"],
                        "listing_price": v["listing_price"]
                    }
                    for v in vars_list
                ]
            }
            all_items.append(MenuItem.model_validate(item_dict))
            
        if all_items:
            return MenuExtraction(currency="INR", items=all_items, document_notes=["Directly parsed from cells"])
            
    except Exception as e:
        print(f"Direct Excel extraction fallback: {e}")
        
    return None


def extract_directly_from_csv(path: Path) -> Optional[MenuExtraction]:
    try:
        import csv
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            rows = list(reader)
        if not rows:
            return None
            
        header_row_idx = -1
        col_map = {} # maps attributes to cell indices
        
        for idx, r in enumerate(rows[:20]): # Check first 20 rows
            cells = [str(c).strip().lower() for c in r if c is not None]
            has_name = any("name" in c or "product" in c or "item" in c for c in cells)
            has_price = any("price" in c or "rate" in c or "mrp" in c or "cost" in c for c in cells)
            
            if has_name and has_price:
                header_row_idx = idx
                for c_idx, val in enumerate(r):
                    if val is None:
                        continue
                    val_str = str(val).strip().lower()
                    if "category" in val_str or "section" in val_str or "group" in val_str:
                        col_map["category"] = c_idx
                    elif "name" in val_str or "product" in val_str or "item" in val_str:
                        col_map["product_name"] = c_idx
                    elif "price" in val_str or "rate" in val_str or "mrp" in val_str or "cost" in val_str:
                        col_map["price"] = c_idx
                    elif "desc" in val_str or "detail" in val_str or "info" in val_str:
                        col_map["description"] = c_idx
                break
                
        if header_row_idx != -1 and "product_name" in col_map and "price" in col_map:
            category_idx = col_map.get("category", -1)
            name_idx = col_map["product_name"]
            price_idx = col_map["price"]
            desc_idx = col_map.get("description", -1)
            
            all_items = []
            current_category = "Uncategorized"
            for r in rows[header_row_idx + 1:]:
                if not any(c is not None and str(c).strip() != "" for c in r):
                    continue
                    
                name_val = r[name_idx] if name_idx < len(r) else None
                if name_val is None or str(name_val).strip() == "" or str(name_val).strip().lower() in ["item name", "product name", "name", "none"]:
                    continue
                    
                if category_idx != -1 and category_idx < len(r) and r[category_idx] is not None:
                    cat_val = str(r[category_idx]).strip()
                    if cat_val:
                        current_category = cat_val
                        
                price_val = r[price_idx] if price_idx < len(r) else None
                price_str = str(price_val).strip() if price_val is not None else ""
                
                # Check cells for inline # names in this row
                hash_cell_names = []
                for cell_val in r:
                    if cell_val is not None and "#" in str(cell_val):
                        cell_str = str(cell_val).strip()
                        if cell_str != price_str and not re.search(r'^\d*(?:\.\d+)?(?:#\d*(?:\.\d+)?)*$', cell_str):
                            hash_cell_names = [x.strip() for x in cell_str.split("#") if x.strip()]
                            break
                            
                row_variations = []
                if "#" in price_str:
                    prices_split = [x.strip() for x in price_str.split("#") if x.strip()]
                    if hash_cell_names and len(hash_cell_names) == len(prices_split):
                        v_names = hash_cell_names
                    else:
                        if len(prices_split) == 2:
                            v_names = ["Half", "Full"]
                        elif len(prices_split) == 3:
                            v_names = ["Small", "Medium", "Large"]
                        else:
                            v_names = [f"Size {k+1}" for k in range(len(prices_split))]
                    
                    for k, p_str in enumerate(prices_split):
                        try:
                            p_val = float(p_str)
                        except ValueError:
                            p_val = 0.0
                        row_variations.append({"name": v_names[k], "price": p_val, "listing_price": p_val})
                else:
                    try:
                        price_num = float(price_val) if price_val is not None and price_str != "" else 0.0
                    except ValueError:
                        price_num = 0.0
                    row_variations.append({"name": "", "price": price_num, "listing_price": price_num})
                    
                desc_val = str(r[desc_idx]).strip() if desc_idx != -1 and desc_idx < len(r) and r[desc_idx] is not None else ""
                
                item_dict = {
                    "category": current_category,
                    "product_name": str(name_val).strip(),
                    "description": desc_val,
                    "dietary_tag": "",
                    "confidence": 1.0,
                    "source_text": f"Direct CSV row import",
                    "variations": row_variations
                }
                
                all_items.append(MenuItem.model_validate(item_dict))
                
            if all_items:
                return MenuExtraction(currency="INR", items=all_items, document_notes=["Directly parsed from CSV cells"])
    except Exception as e:
        print(f"Direct CSV extraction fallback: {e}")
    return None


def extract_text_from_image_via_ocr(image_path: Path) -> str:
    # 1. Try RapidOCR first
    try:
        from rapidocr_onnxruntime import RapidOCR
        engine = RapidOCR()
        result, elapse = engine(str(image_path))
        if result:
            lines_data = []
            for item in result:
                box = item[0]
                text = item[1]
                
                y0 = box[0][1]
                x0 = box[0][0]
                lines_data.append({"text": text, "x": x0, "y": y0})
                
            lines_data = sorted(lines_data, key=lambda it: it["y"])
            grouped_lines = []
            if lines_data:
                current_line = [lines_data[0]]
                for item in lines_data[1:]:
                    # Group items within a 20-pixel vertical tolerance
                    if abs(item["y"] - current_line[0]["y"]) < 20:
                        current_line.append(item)
                    else:
                        grouped_lines.append(current_line)
                        current_line = [item]
                grouped_lines.append(current_line)
                
            final_lines = []
            for grp in grouped_lines:
                grp_sorted = sorted(grp, key=lambda it: it["x"])
                line_str = " ".join([item["text"] for item in grp_sorted])
                final_lines.append(line_str.strip())
                
            text_out = "\n".join(final_lines)
            if text_out.strip():
                return text_out
    except Exception as e:
        print(f"Error in local RapidOCR: {e}. Falling back to Tesseract OCR...")

    # 2. Fallback to Tesseract OCR (standard system binary)
    try:
        tess_paths = [
            "tesseract",
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
        ]
        
        tess_bin = None
        for path in tess_paths:
            if shutil.which(path):
                tess_bin = path
                break
            elif Path(path).exists():
                tess_bin = path
                break
                
        if not tess_bin:
            print("Tesseract binary not found in PATH or standard Program Files locations.")
            return ""
            
        print(f"Executing Tesseract OCR: {tess_bin} on {image_path}")
        res = subprocess.run(
            [tess_bin, str(image_path), "stdout"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            check=True
        )
        tess_text = res.stdout.strip()
        if tess_text:
            print(f"Tesseract OCR extracted successfully ({len(tess_text)} chars).")
            return tess_text
    except Exception as ex:
        print(f"Error in local Tesseract OCR: {ex}")
        
    return ""


def extract_menu_from_file(path: str | Path, engine: str = "auto", api_key: Optional[str] = None) -> MenuExtraction:
    path = Path(path)
    suffix = path.suffix.lower()
    engine = str(engine).lower().strip()

    if suffix not in SUPPORTED_EXTS:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Supported: images {sorted(SUPPORTED_IMAGE_EXTS)}, "
            f"PDF, Excel {sorted(SUPPORTED_EXCEL_EXTS)}, Word {sorted(SUPPORTED_WORD_EXTS)}."
        )

    # 1. Optimize Excel/CSV Tabular files: try parser direct bypass to save massive time
    if suffix in SUPPORTED_EXCEL_EXTS:
        if suffix == ".csv":
            direct = extract_directly_from_csv(path)
        else:
            direct = extract_directly_from_excel(path)
        if direct and direct.items:
            print(f"Successfully bypassed AI for direct cell extraction: {len(direct.items)} products.")
            return direct

    # 2. Image menu: Try Ollama/Gemini vision model directly first!
    if suffix in SUPPORTED_IMAGE_EXTS:
        temp_path = Path(tempfile.mkdtemp(prefix="menu_img_")) / "input.jpg"
        _compress_image_for_vision(path, temp_path)
        
        use_gemini = False
        use_ollama = False
        use_heuristics = False
        
        gemini_key = api_key or os.getenv("GEMINI_API_KEY")
        
        if engine == "gemini":
            use_gemini = True
        elif engine == "ollama":
            use_ollama = True
        elif engine == "heuristics":
            use_heuristics = True
        else: # "auto"
            if gemini_key:
                use_gemini = True
            else:
                use_ollama = True
                
        if use_gemini:
            print("Force/Auto Gemini vision model for image extraction...")
            try:
                return extract_from_images_with_gemini([temp_path], api_key=gemini_key)
            except Exception as e:
                if engine == "gemini":
                    print(f"Forced Gemini extraction failed (429/quota/network): {e}. Gracefully falling back to local Heuristics OCR...")
                    use_heuristics = True
                    gemini_failed_note = f"[IMPORTANT] Forced Gemini extraction failed: {e}. Automatically fell back to local OCR + Heuristics."
                else:
                    print(f"Gemini Vision failed, trying Ollama: {e}")
                    use_ollama = True
                
        if use_ollama:
            print("Force/Auto Ollama vision model for image extraction...")
            try:
                return extract_from_image_with_ollama(temp_path, api_key=gemini_key, bypass_to_gemini=(engine != "ollama"))
            except Exception as e:
                if engine == "ollama":
                    print(f"Forced Ollama extraction failed: {e}. Gracefully falling back to local Heuristics OCR...")
                    use_heuristics = True
                    ollama_failed_note = f"[IMPORTANT] Forced Ollama extraction failed: {e}. Automatically fell back to local OCR + Heuristics."
                else:
                    print(f"Ollama Vision failed: {e}. Falling back to local RapidOCR + Heuristics...")
                    use_heuristics = True
                
        if use_heuristics:
            print("Using local RapidOCR + Heuristics for image extraction...")
            ocr_text = extract_text_from_image_via_ocr(temp_path)
            extracted = MenuExtraction(currency="INR", items=[], document_notes=["Heuristic extraction failed to find items."])
            if ocr_text.strip():
                heur = parse_menu_text_heuristically(ocr_text)
                if heur and heur.items:
                    print(f"Processed via local OCR + Heuristics: {len(heur.items)} items.")
                    extracted = heur
            if 'gemini_failed_note' in locals():
                extracted.document_notes.append(gemini_failed_note)
            if 'ollama_failed_note' in locals():
                extracted.document_notes.append(ollama_failed_note)
            return extracted

    # 3. PDF: use embedded text when available; scanned PDF fallback goes page-by-page.
    if suffix == ".pdf":
        text = extract_plain_text_from_file(path)
        is_scanned = len(text.strip()) <= 300
        
        use_gemini = False
        use_ollama = False
        use_heuristics = False
        
        gemini_key = api_key or os.getenv("GEMINI_API_KEY")
        
        if engine == "gemini":
            use_gemini = True
        elif engine == "ollama":
            use_ollama = True
        elif engine == "heuristics":
            use_heuristics = True
        else: # "auto"
            if gemini_key:
                use_gemini = True
            else:
                use_ollama = True
                
        if not is_scanned:
            # Native text PDF
            if use_gemini:
                try:
                    print("Using Gemini for text PDF extraction...")
                    return extract_from_text_with_gemini(text, api_key=gemini_key)
                except Exception as e:
                    if engine == "gemini":
                        print(f"Forced Gemini text extraction failed: {e}. Gracefully falling back to local Heuristics...")
                        use_heuristics = True
                        gemini_failed_note = f"[IMPORTANT] Forced Gemini extraction failed: {e}. Automatically fell back to local Heuristics."
                    else:
                        print(f"Gemini text extraction failed: {e}. Trying Ollama...")
                        use_ollama = True
            if use_ollama:
                try:
                    print("Using Ollama for text PDF extraction...")
                    chunks = _chunk_text_by_lines(text, max_chars=3000)[:3]
                    extractions = [extract_from_text_with_ollama(chunk, api_key=gemini_key, bypass_to_gemini=(engine != "ollama")) for chunk in chunks]
                    return merge_extractions(extractions)
                except Exception as e:
                    if engine == "ollama":
                        print(f"Forced Ollama text extraction failed: {e}. Gracefully falling back to local Heuristics...")
                        use_heuristics = True
                        ollama_failed_note = f"[IMPORTANT] Forced Ollama extraction failed: {e}. Automatically fell back to local Heuristics."
                    else:
                        print(f"Ollama text extraction failed: {e}. Falling back to Heuristics...")
                        use_heuristics = True
            if use_heuristics:
                print("Using local Heuristics for PDF text extraction...")
                extracted = MenuExtraction(currency="INR", items=[], document_notes=["Local heuristics failed to extract menu items from text."])
                heur = parse_menu_text_heuristically(text)
                if heur and heur.items:
                    extracted = heur
                if 'gemini_failed_note' in locals():
                    extracted.document_notes.append(gemini_failed_note)
                if 'ollama_failed_note' in locals():
                    extracted.document_notes.append(ollama_failed_note)
                return extracted
        else:
            # Scanned PDF: page-by-page vision
            image_paths = _pdf_pages_to_images(path, max_pages=10)
            if use_gemini:
                try:
                    print("Using Gemini for scanned PDF vision extraction...")
                    return extract_from_images_with_gemini(image_paths, api_key=gemini_key)
                except Exception as e:
                    if engine == "gemini":
                        print(f"Forced Gemini scanned PDF extraction failed: {e}. Gracefully falling back to local OCR + Heuristics...")
                        use_heuristics = True
                        gemini_failed_note = f"[IMPORTANT] Forced Gemini extraction failed: {e}. Automatically fell back to local OCR + Heuristics."
                    else:
                        print(f"Gemini scanned PDF vision failed: {e}. Trying Ollama...")
                        use_ollama = True
            if use_ollama:
                try:
                    print("Using Ollama for scanned PDF vision extraction...")
                    extractions = [extract_from_image_with_ollama(img_path, api_key=gemini_key, bypass_to_gemini=(engine != "ollama")) for img_path in image_paths[:3]]
                    return merge_extractions(extractions)
                except Exception as e:
                    if engine == "ollama":
                        print(f"Forced Ollama scanned PDF extraction failed: {e}. Gracefully falling back to local OCR + Heuristics...")
                        use_heuristics = True
                        ollama_failed_note = f"[IMPORTANT] Forced Ollama extraction failed: {e}. Automatically fell back to local OCR + Heuristics."
                    else:
                        print(f"Ollama scanned PDF vision failed: {e}. Falling back to Local OCR...")
                        use_heuristics = True
            if use_heuristics:
                print("Using local OCR + Heuristics for scanned PDF...")
                ocr_texts = []
                extracted = MenuExtraction(currency="INR", items=[], document_notes=["Local OCR + Heuristics failed to extract menu items from scanned PDF."])
                for img_path in image_paths:
                    ocr_texts.append(extract_text_from_image_via_ocr(img_path))
                combined_text = "\n".join(ocr_texts)
                if combined_text.strip():
                    heur = parse_menu_text_heuristically(combined_text)
                    if heur and heur.items:
                        extracted = heur
                if 'gemini_failed_note' in locals():
                    extracted.document_notes.append(gemini_failed_note)
                if 'ollama_failed_note' in locals():
                    extracted.document_notes.append(ollama_failed_note)
                return extracted

    # 4. Word/Text: convert to clean text and send chunked to AI text model.
    text = extract_plain_text_from_file(path)
    if not text.strip():
        raise ValueError(f"Could not read menu content from: {path.name}")
        
    use_gemini = False
    use_ollama = False
    use_heuristics = False
    
    gemini_key = api_key or os.getenv("GEMINI_API_KEY")
    
    if engine == "gemini":
        use_gemini = True
    elif engine == "ollama":
        use_ollama = True
    elif engine == "heuristics":
        use_heuristics = True
    else: # "auto"
        if gemini_key:
            use_gemini = True
        else:
            use_ollama = True
            
    if use_gemini:
        try:
            print("Using Gemini for Word/Text document extraction...")
            return extract_from_text_with_gemini(text, api_key=gemini_key)
        except Exception as e:
            if engine == "gemini":
                print(f"Forced Gemini Word/Text extraction failed: {e}. Gracefully falling back to local Heuristics...")
                use_heuristics = True
                gemini_failed_note = f"[IMPORTANT] Forced Gemini extraction failed: {e}. Automatically fell back to local Heuristics."
            else:
                print(f"Gemini Word/Text extraction failed: {e}. Trying Ollama...")
                use_ollama = True
    if use_ollama:
        try:
            print("Using Ollama for Word/Text document extraction...")
            chunks = _chunk_text_by_lines(text, max_chars=3000)[:3]
            extractions = [extract_from_text_with_ollama(chunk, api_key=gemini_key, bypass_to_gemini=(engine != "ollama")) for chunk in chunks]
            return merge_extractions(extractions)
        except Exception as e:
            if engine == "ollama":
                print(f"Forced Ollama Word/Text extraction failed: {e}. Gracefully falling back to local Heuristics...")
                use_heuristics = True
                ollama_failed_note = f"[IMPORTANT] Forced Ollama extraction failed: {e}. Automatically fell back to local Heuristics."
            else:
                print(f"Ollama Word/Text extraction failed: {e}. Falling back to Heuristics...")
                use_heuristics = True
    if use_heuristics:
        print("Using local Heuristics for Word/Text extraction...")
        extracted = MenuExtraction(currency="INR", items=[], document_notes=["Local heuristics failed to extract menu items from Word/Text file."])
        heur = parse_menu_text_heuristically(text)
        if heur and heur.items:
            extracted = heur
        if 'gemini_failed_note' in locals():
            extracted.document_notes.append(gemini_failed_note)
        if 'ollama_failed_note' in locals():
            extracted.document_notes.append(ollama_failed_note)
        return extracted

